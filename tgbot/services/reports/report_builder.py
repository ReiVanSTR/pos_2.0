from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from tgbot.models import ObjectId, Session
from tgbot.models import Shift as ShiftModel
from .session_dataclasses import SessionReportData, EmployerSellings, Shift, SessionData

from enum import Enum
from typing import Union
from datetime import datetime
from pytz import timezone

tzinfo = timezone("Europe/Warsaw")

class DocumentBuilder():
    path: str

    def __init__(self, path: str = "./") -> None:
        self.path = path
        self.document = Document()

    def add_title(self, text):
        heading = self.document.add_heading(level=0)
        run = heading.add_run()
        run.text = text
        run.bold = True
        run.font.size = Pt(32)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_heading(self, text, aligment: WD_PARAGRAPH_ALIGNMENT, color: RGBColor = RGBColor(0,0,0)):
        heading = self.document.add_heading(level = 2)
        run = heading.add_run()
        run.text = text
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = color
        heading.alignment = aligment
        p = self.document.add_paragraph()
        x = p.add_run("text")


    def add_section_heading(self, text, aligment: WD_PARAGRAPH_ALIGNMENT, color: RGBColor = RGBColor(0,0,0)):
        heading = self.document.add_heading(level = 4)
        run = heading.add_run()
        run.text = text
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color
        heading.alignment = aligment

    def add_employer_entry(self, employer_sellings: EmployerSellings, employer_shift: Shift):
        if employer_shift:
            shift_text = f"{employer_shift.total_hours}h {employer_shift.total_minutes} min"
        else:
            shift_text = "Brak godzin pracy"

        self.add_section_heading(text = f"{employer_sellings.employer_name} ({shift_text}) \n", aligment = WD_PARAGRAPH_ALIGNMENT.LEFT)
        paragraf = self.document.add_paragraph()
        paragraf.add_run(text = "Sprzedaż na kartę: \n").bold = True
        if len(employer_sellings.bills_by_card) < 1:
            paragraf.add_run(text = f"\t Brak realizowanych transakcji\n") 
        for selling in employer_sellings.bills_by_card:
            paragraf.add_run(text = f"\t - {selling.bill_name} | {selling.bill_cost} pln. Zamówień: {selling.orders_count}\n").italic = True
        paragraf.add_run(text = f"Łączna sprzedaż na kartę: {employer_sellings.card_total}\n\n").font.cs_bold = True

        paragraf.add_run(text = "Sprzedaż gotówką:\r").bold = True
        if len(employer_sellings.bills_by_cash) < 1:
            paragraf.add_run(text = f"\t Brak realizowanych transakcji\n") 
        for selling in employer_sellings.bills_by_cash:
            paragraf.add_run(text = f"\t - {selling.bill_name} | {selling.bill_cost} pln. Zamówień: {selling.orders_count}\n").italic = True
        paragraf.add_run(text = f"Łączna sprzedaż gotówką: {employer_sellings.cash_total}\n\n").font.cs_bold = True


        paragraf.add_run(text = "Sprzedaż wewnętrzna (Szef):\n").bold = True
        if len(employer_sellings.chief) < 1:
            paragraf.add_run(text = f"\t Brak realizowanych transakcji\n") 
        for selling in employer_sellings.chief:
            paragraf.add_run(text = f"\t - {selling.bill_name} | {selling.bill_cost} pln. Zamówień: {selling.orders_count}\n").italic = True
        paragraf.add_run(text = f"Łączna sprzedaż wewnętrzna: {employer_sellings.chief_total}\n\n").font.cs_bold = True

        paragraf.add_run(text = f"Łączna sprzedaż: {employer_sellings.card_total + employer_sellings.cash_total} pln. (karta + gotówka)\n")
        paragraf.add_run(text = f"Ilość realizowanych zamówień: {employer_sellings.total_orders}\n")
        paragraf.add_run(text = f"Ilość fiskalizowanych: {employer_sellings.total_sellings}\n")

    def add_session_info_entry(self, session_data: SessionData):
        paragraf = self.document.add_paragraph()
        paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraf.add_run(text = f"""
                        Rozpoczęcie sprzedaży: {session_data.opened_by.timestamp} ({session_data.opened_by.user_data.username}) 
                        Koniec sprzedaży: {session_data.closed_by.timestamp} ({session_data.closed_by.user_data.username}) 
                        Łączny czas: {session_data.session_active_time.hours} h {session_data.session_active_time.minutes} min.
        """).font.bold = True
        
    def add_session_tabacco_data(self, session_data: SessionReportData):
        self.add_section_heading(text = f"Uzyty tytoń", aligment = WD_PARAGRAPH_ALIGNMENT.CENTER)
        paragraf = self.document.add_paragraph()
        paragraf.add_run(text = f"\n".join([
            f"{tabacco.label} - {tabacco.total_used}g" for tabacco in session_data.tabacco_data
        ]))

    def add_sesion_total_entry(self, session_data: SessionReportData):
        self.add_section_heading(text = f"Total", aligment = WD_PARAGRAPH_ALIGNMENT.CENTER)
        paragraf = self.document.add_paragraph()
        paragraf.add_run(text = f"""
                        Karta: {session_data.total_selling_by_card}
                        Gotówka: {session_data.total_selling_cash}
                        Szef: {session_data.total_selling_chief}
                        Użyty tytoń: {session_data.total_tabacco}g
        """)
        
    def add_employer_sellings_table(self, sellings_data):
        table = self.document.add_table(rows = 1, cols = 7)
        header_cells_names = ["Data", "Początek Zmiany", "Koniec Zmiany", "Godziny", "Minuty", "Sprzedaż", "Premia od sprzedaży"]
        for cell in range(len(table.rows[0].cells)):
            table.rows[0].cells[cell].text = header_cells_names[cell]
        
        total_prem = 0
        
        for selling in sellings_data.get("shifts"):
            prem = (selling.get("sellings", 0) * 5) * (2 if selling.get("sellings", 0) >= 10 else 1)
            total_prem += prem
            row_cells = table.add_row().cells
            row_cells[0].text = selling.get("date")
            row_cells[1].text = selling.get("start_time").strftime("%d-%m-%Y %H:%M")
            row_cells[2].text = selling.get("end_time").strftime("%d-%m-%Y %H:%M")
            row_cells[3].text = str(selling.get("hours"))
            row_cells[4].text = str(selling.get("minutes"))
            row_cells[5].text = str(selling.get("sellings", 0))
            row_cells[6].text = str(prem)
            
        paragraf = self.document.add_paragraph()
        paragraf.add_run("\n")
        self.add_section_heading(text = f"Total", aligment = WD_PARAGRAPH_ALIGNMENT.CENTER)
        paragraf = self.document.add_paragraph()
        paragraf.add_run(text = f"""
                        Godziny: {sellings_data.get("total_hours")}
                        Minuty: {sellings_data.get("total_minutes")}
                        Ilość zmian: {len(sellings_data.get("shifts"))}
                        Rozliczenie według godzin: False
                        Premia za zmianę: 150 PLN
                        Łączna premia zmianowa: {len(sellings_data.get("shifts")) * 150}
                        Premia od sprzedaży: {total_prem}
        """)
        
        
        
    def save(self, path: str = "./", document_name: str = "document"):
        if path:
            self.path = path

        self.document.save(f"{path}/{document_name}.docx")
        self.document = Document()

class ReportType(Enum):
    CHANGE = "CHANGE_REPORT"
    PERIODIC = "PERIODIC_REPORT"




class Report():

    def __init__(self) -> None:
        self.builder = DocumentBuilder()
        self.default_route = "./reports"

    async def generate_change_report(self, session_id: Union[ObjectId, str], user_name: str, filename: str = None, session_data = None):
        _session_data: SessionReportData = await Session.generate_report_data(session_id = session_id) if not session_data else session_data
        _generating_date = datetime.now(tz=tzinfo).strftime("%d-%m-%Y %H:%M") 

        self.builder.add_title(f"Raport zmianowy nr {_session_data.session_data.session_id.__str__()[8:15]}")
        self.builder.add_section_heading(text = f"Data wygenerowania raportu: {_generating_date}. \n Wygenerowany przez: {user_name} \n", aligment=WD_PARAGRAPH_ALIGNMENT.CENTER)

        self.builder.add_session_info_entry(_session_data.session_data)

        for employer_data in _session_data.employer_sellings:
            shift = _session_data.find_shift(employer_data.employer_name)

            self.builder.add_employer_entry(employer_data, shift)
            self.builder.document.add_page_break()
        self.builder.add_session_tabacco_data(_session_data)
        self.builder.add_sesion_total_entry(_session_data)

        document_name = f"raport_zmianowy_{datetime.now(tz=tzinfo).strftime('%d-%m-%Y')}"

        if filename:
            document_name = filename

        self.builder.save(path = self.default_route, document_name = document_name)
        
    async def generate_employer_report(self, user_id: int, from_date: datetime, to_date: datetime ,user_name: str, filename: str = None):
        _employer_data = await ShiftModel.generate_total_report_data(user_id, from_date, to_date)
        if not _employer_data:
            return 
        _employer_data = _employer_data[0]
        _generating_date = datetime.now(tz=tzinfo).strftime("%d-%m-%Y %H:%M") 
        
        
        self.builder.add_title(f"Rozliczenie wypłaty {_employer_data.get('username')}")
        self.builder.add_section_heading(text = f"Data wygenerowania raportu: {_generating_date}. \n Wygenerowany przez: {user_name} \n", aligment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        
        
        self.builder.add_employer_sellings_table(_employer_data)
        
        document_name = f"paski_rozliczeniowe_{_employer_data.get('username')}_{from_date.strftime('%m-%Y')}"

        if filename:
            document_name = filename

        self.builder.save(path = self.default_route, document_name = document_name)